import docx, os

AOD_FILEPATH = 'S:\\AOD\\2017'
GEN_FILEPATH = 'S:\\Deputy Clerks'
PATH = 'S:\\AOD\\AOD_Test\\'
TEMPLATE_PATH = "S:\\Letter_Writer\\Templates\\"
SAVE_PATH = "S:\\Letter_Writer\\"

TEMPLATE = TEMPLATE_PATH + 'Template.docx'
AFFIANT_TEMPLATE = TEMPLATE_PATH + 'Affiant_Template.docx'
JUDGE_TEMPLATE = TEMPLATE_PATH + 'Judge_Template.docx'
GEN_TEMPLATE = TEMPLATE_PATH + 'Gen_Template.docx'
NOTFILED_TEMPLATE = TEMPLATE_PATH + 'NotFiled_Template.docx'
NOCASE_TEMPLATE = TEMPLATE_PATH + 'NoCase_Template.docx'
NOFORMS_TEMPLATE = TEMPLATE_PATH + 'NoForms_Template.docx'
LATEJUR_TEMPLATE = TEMPLATE_PATH + 'LateJur_Template.docx'
LATEJUR_DAF_TEMPLATE = TEMPLATE_PATH + 'LateJurDelayedAppeal_Template.docx'
JUR_MISSINGDOCS_TEMPLATE = TEMPLATE_PATH + 'JurMissingDocs_Template.docx'


class Letter(object):
    def __init__(self, fields):
        self.path = SAVE_PATH
        self.fields = fields
        self.letter = docx.Document(TEMPLATE)
        for paragraph in self.letter.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None

    def create_letter(self):
        """self.template is not initialized for superclass, in the
        subclasses it is assigned the relevant template."""
        self.master_template = docx.Document(self.template)
        self.string = ""
        for paragraph in self.master_template.paragraphs:
            self.string = self.string + '\n' + paragraph.text
        self.newstring = self.string.format(**self.fields)
        self.letter.add_paragraph(self.newstring)
        self.letter.save(self.docname)
        self.open_letter()

    def open_letter(self):
        os.startfile(self.docname)


class AOD_AffiantLetter(Letter):
    def __init__(self, fields):
        Letter.__init__(self, fields)
        self.template = AFFIANT_TEMPLATE
        self.docname = self.path + self.fields["affiant_last_name"] + ', ' +\
            self.fields["affiant_first_name"] + " (" + self.fields["date"] + ")" +\
            '.docx'

    def create_letter(self):
        """Need to modify this to put rejection reasons in an indented
        bullet list format. String, then Docx Add Paragraph, then String.
        Or if fix is done elsewhere eliminate this method."""
        self.master_template = docx.Document(self.template)
        self.string = ""
        for paragraph in self.master_template.paragraphs:
            self.string = self.string + '\n' + paragraph.text
        self.newstring = self.string.format(**self.fields)
        self.letter.add_paragraph(self.newstring)
        self.letter.save(self.docname)
        self.open_letter()


class AOD_JudgeLetter(Letter):
    def __init__(self, fields):
        Letter.__init__(self, fields)
        self.template = JUDGE_TEMPLATE
        self.docname = self.path + self.fields["judge_last_name"] + ', ' +\
            self.fields["judge_first_name"] + " (" + self.fields["date"] + ")" +\
            '.docx'


class GEN_Letter(Letter):
    def __init__(self, fields):
        Letter.__init__(self, fields)
        self.template = GEN_Letter.get_template()
        self.docname = self.path + self.fields["last_name"] + ', ' +\
            self.fields["first_name"] + " (" + self.fields["date"] + ")" +\
            '.docx'

    @classmethod
    def get_template(cls):
        return GEN_TEMPLATE

    @classmethod
    def return_preview(cls):
        preview_template = docx.Document(cls.get_template())
        string = ""
        for paragraph in preview_template.paragraphs:
            string = string + '\n' + paragraph.text
        new = string.split(':', maxsplit=1)
        new = new[1].split('\n', maxsplit=2)
        return new[2:]


class GEN_NotFiledLetter(GEN_Letter):
    def __init__(self, fields):
        GEN_Letter.__init__(self, fields)
        self.template = NOTFILED_TEMPLATE

    @classmethod
    def get_template(cls):
        return NOTFILED_TEMPLATE


class GEN_NoCaseLetter(GEN_Letter):
    def __init__(self, fields):
        GEN_Letter.__init__(self, fields)
        self.template = NOCASE_TEMPLATE

    @classmethod
    def get_template(cls):
        return NOCASE_TEMPLATE


class GEN_NoFormsLetter(GEN_Letter):
    def __init__(self, fields):
        GEN_Letter.__init__(self, fields)
        self.template = NOFORMS_TEMPLATE

    @classmethod
    def get_template(cls):
        return NOFORMS_TEMPLATE


class GEN_LateJurLetter(GEN_Letter):
    def __init__(self, fields):
        GEN_Letter.__init__(self, fields)
        self.template = LATEJUR_TEMPLATE

    @classmethod
    def get_template(cls):
        return LATEJUR_TEMPLATE


class GEN_LateJurDelayedAppealLetter(GEN_Letter):
    def __init__(self, fields):
        GEN_Letter.__init__(self, fields)
        self.template = LATEJUR_DAF_TEMPLATE

    @classmethod
    def get_template(cls):
        return LATEJUR_DAF_TEMPLATE


class GEN_TimelyJurMissingDocsLetter(GEN_Letter):
    def __init__(self, fields):
        GEN_Letter.__init__(self, fields)
        self.template = JUR_MISSINGDOCS_TEMPLATE

    @classmethod
    def get_template(cls):
        return JUR_MISSINGDOCS_TEMPLATE
